import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def _option_text(opts: list[dict], key: str | None) -> str:
    if not key:
        return ""
    for o in opts:
        if o.get("key") == key:
            return o.get("text", "")
    return ""


def build_docx(
    user_email: str,
    channel_title: str,
    mcqs: list[dict],
    user_answers: dict[str, str] | None = None,
    subtitle: str | None = None,
) -> bytes:
    """Answer-key DOCX.

    Answer precedence per question:
      1. User's selected answer from the quiz page (preferred — the client asked for this).
      2. Channel's correct answer (if Telegram revealed it).
      3. "Not attempted" placeholder.
    """
    user_answers = user_answers or {}

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TeleMCQ Export")
    run.bold = True
    run.font.size = Pt(24)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = subtitle_para.add_run(subtitle or "All MCQs")
    st.italic = True
    st.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Channel: {channel_title}\n").bold = True
    meta.add_run(f"User: {user_email}\n")
    meta.add_run(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}\n")
    meta.add_run(f"Total MCQs: {len(mcqs)}")

    doc.add_page_break()

    current_category: object | str = object()
    for n, m in enumerate(mcqs, 1):
        cat = m.get("category") or "Uncategorised"
        if cat != current_category:
            current_category = cat
            h = doc.add_paragraph()
            hr = h.add_run(cat)
            hr.bold = True
            hr.font.size = Pt(14)
            hr.font.color.rgb = RGBColor(0x00, 0x6E, 0xC1)

        q = doc.add_paragraph()
        qr = q.add_run(f"Q{n}. {m['question']}")
        qr.bold = True
        qr.font.size = Pt(12)

        user_sel = user_answers.get(m["id"])
        channel_correct = m.get("correct_answer")
        # Priority: the *actual* correct answer (if Telegram revealed it) wins over
        # the user's selection — a wrong attempt shouldn't be presented as the answer.
        answer_key = channel_correct or user_sel
        used_user_answer = answer_key is not None and answer_key == user_sel and not channel_correct

        for opt in m["options"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(f"{opt['key']}) {opt['text']}")
            if answer_key and opt["key"] == answer_key:
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

        ans_line = doc.add_paragraph()
        if answer_key:
            label = "Your answer" if used_user_answer else "Answer"
            ar = ans_line.add_run(
                f"{label}: {answer_key}) {_option_text(m['options'], answer_key)}"
            )
            ar.bold = True
            ar.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        else:
            ar = ans_line.add_run("Answer: (not available)")
            ar.italic = True
            ar.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # If we showed the correct answer AND the user had attempted (and was wrong), note it.
        if channel_correct and user_sel and user_sel != channel_correct:
            note_line = doc.add_paragraph()
            nr = note_line.add_run(f"(You chose: {user_sel})")
            nr.italic = True
            nr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
